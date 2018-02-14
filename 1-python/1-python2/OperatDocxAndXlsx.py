#!/usr/bin/env python
# coding=utf-8
import xlwt
'''
reference：http://my.oschina.net/alazyer/blog/223354
'''

def saveDoc():
    from docx import Document
    from docx.shared import Inches

    document = Document()
    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')
    document.add_paragraph('first item in unordered list', style='ListBullet')
    document.add_paragraph('first item in ordered list', style='ListNumber')

    document.add_picture('/home/alex/BigFolder/android-sdk/platforms/android-14/data/res/drawable-hdpi/ic_dialog_dialer.png', width=Inches(1.25))

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
#    for item in recordset:
#        row_cells = table.add_row().cells
#        row_cells[0].text = str(item.qty)
#        row_cells[1].text = str(item.id)
#        row_cells[2].text = item.desc

    document.add_page_break()

    document.save('demo.docx')

'''
workbook = xlwt.Workbook(encoding='utf-8') #创建一个文件
booksheet = workbook.add_sheet('Sheet 1') #增加一个表单
row = booksheet.row(i) #增加一行
col = booksheet.col(j) #增加一列
#为某个单元格赋值
booksheet.write(i, j, 'Test')
row.write(j, 'Test')
col.write(i, 'Test')
注意: 在add_sheet时, 置参数cell_overwrite_ok=True, 可以覆盖原单元格中数据。
cell_overwrite_ok默认为False, 覆盖的话, 会抛出异常.
'''
def saveXlsx():
    workbook = xlwt.Workbook(encoding='utf-8')#创建一个文件
    booksheet = workbook.add_sheet('Sheet 1', cell_overwrite_ok=True)#增加一个表单
    workbook.add_sheet('Sheet 2')
    DATA = (('学号', '姓名', '年龄', '性别', '成绩'),
            (1001, 'AAAA', 23, '男', 98),
            (1002, 'BBBB', 21, '女', 90),
            (1003, 'CCCC', 24, '女', 100),
            (1004, 'DDDD', 22, '女', 86),
            (1005, 'EEEE', 25, '女', 88),)

    for i, row in enumerate(DATA):
        for j, col in enumerate(row):
            booksheet.write(i, j, col)
    booksheet.col(0).width = 10
    workbook.save('成绩单.xlsx')

'''
workbook = xlrd.open_workbook('成绩单.xls') #获取一个文件
workbook.nsheets #获取文件中表单数量
workbook.sheets()[i], #获取一个表单
workbook.sheet_by_index(i)
workbook.sheet_by_name(u'Sheet1')
sheet.nrows, sheet.ncols #获取行，列数
sheet.row(i), sheet.col(j) #获取整行，列数据
#获取某个单元格数据
sheet.cell(i, j).value
sheet.row(i)[j].value
sheet.col(j)[i].value
'''
def readXlsx():
    import xlrd
    workbook = xlrd.open_workbook('成绩单.xlsx')
    print "There are {} sheets in the workbook".format(workbook.nsheets)
    for booksheet in workbook.sheets():
        print booksheet.name
        for row in xrange(booksheet.nrows):
            for col in xrange(booksheet.ncols):
                print xlrd.cellname(row, col)
                print booksheet.cell(row, col).value

if __name__ == '__main__':
    saveXlsx()
    readXlsx()
    print "Successed!"